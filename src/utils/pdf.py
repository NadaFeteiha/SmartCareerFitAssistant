"""Professional PDF and Word generation utilities for resumes and cover letters."""

from io import BytesIO
from datetime import datetime
import logging
import re

logger = logging.getLogger(__name__)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

# ── Page geometry ─────────────────────────────────────────────────────────────
PAGE_WIDTH, PAGE_HEIGHT = letter
L_MAR = R_MAR = 0.6 * inch
T_MAR = B_MAR = 0.65 * inch
CONTENT_W = PAGE_WIDTH - L_MAR - R_MAR   # ~7.3 inch

# ── Color palette ─────────────────────────────────────────────────────────────
DARK_GRAY   = colors.HexColor("#0F172A")
ACCENT_LINE = colors.HexColor("#2C3E50")
BODY_TEXT   = colors.HexColor("#1E293B")
LIGHT_GRAY  = colors.HexColor("#475569")
BLUE_ACCENT = colors.HexColor("#3498DB")

SKILL_LABEL_BG   = colors.HexColor("#E0E7FF")
SKILL_ODD_ROW_BG = colors.HexColor("#EEF2FF")
SKILL_BORDER     = colors.HexColor("#CBD5E1")


# ─────────────────────────────────────────────────────────────────────────────
# Typography
# ─────────────────────────────────────────────────────────────────────────────
def _styles() -> dict:
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "Name", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=24, leading=28,
            textColor=DARK_GRAY, alignment=TA_CENTER, spaceAfter=3,
        ),
        # ← NEW: job title line under the name
        "job_title_header": ParagraphStyle(
            "JobTitleHeader", parent=base["Normal"],
            fontName="Helvetica", fontSize=12, leading=16,
            textColor=BLUE_ACCENT, alignment=TA_CENTER, spaceAfter=6,
        ),
        "contact": ParagraphStyle(
            "Contact", parent=base["Normal"],
            fontName="Helvetica", fontSize=9.5, leading=13,
            textColor=BODY_TEXT, alignment=TA_CENTER, spaceAfter=14,
        ),
        "section": ParagraphStyle(
            "SectionHeader", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=11.5, leading=14,
            textColor=DARK_GRAY, spaceBefore=14, spaceAfter=3,
        ),
        # Left side of experience row: Role, Company
        "exp_left": ParagraphStyle(
            "ExpLeft", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=10.5, leading=14,
            textColor=DARK_GRAY,
        ),
        # Right side of experience row: Date
        "exp_right": ParagraphStyle(
            "ExpRight", parent=base["Normal"],
            fontName="Helvetica", fontSize=9.8, leading=14,
            textColor=LIGHT_GRAY, alignment=TA_RIGHT,
        ),
        "bullet": ParagraphStyle(
            "Bullet", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=13.5,
            leftIndent=14, firstLineIndent=-14,
            spaceBefore=1, spaceAfter=3, textColor=BODY_TEXT,
        ),
        "normal": ParagraphStyle(
            "NormalP", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=13.5,
            textColor=BODY_TEXT, spaceAfter=5,
        ),
        "skill_label": ParagraphStyle(
            "SkillLabel", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=10, leading=13,
            textColor=DARK_GRAY,
        ),
        "skill_value": ParagraphStyle(
            "SkillValue", parent=base["Normal"],
            fontName="Helvetica", fontSize=9.5, leading=13,
            textColor=BODY_TEXT,
        ),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def _hr(thickness=0.8, color=ACCENT_LINE, spaceBefore=3, spaceAfter=7) -> HRFlowable:
    return HRFlowable(
        width="100%", thickness=thickness, lineCap="round",
        color=color, spaceBefore=spaceBefore, spaceAfter=spaceAfter,
    )


def _md_to_xml(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*',     r'<i>\1</i>', text)
    text = re.sub(r'&(?!(amp|lt|gt|nbsp|apos|quot);)', r'&amp;', text)
    return text


def _exp_row(role: str, company: str, date: str, location: str, st: dict) -> Table:
    """
    Single experience header row:
      LEFT  — Role (bold) · Company
      RIGHT — Date  (light gray, right-aligned)
              Location (italic, smaller, right-aligned)  [optional]
    """
    left_xml = f"<b>{role}</b>"
    if company:
        left_xml += f",  <font name='Helvetica' color='#1e293b'>{company}</font>"

    right_lines = [f"<b>{date}</b>"] if date else []
    if location:
        right_lines.append(
            f"<font name='Helvetica-Oblique' size='9'>{location}</font>"
        )
    right_xml = "<br/>".join(right_lines)

    t = Table(
        [[Paragraph(left_xml, st["exp_left"]),
          Paragraph(right_xml, st["exp_right"])]],
        colWidths=[CONTENT_W * 0.65, CONTENT_W * 0.35],
        spaceBefore=9, spaceAfter=2,
    )
    t.setStyle(TableStyle([
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING",    (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ("LEFTPADDING",   (0, 0), (0,  0),  0),
        ("RIGHTPADDING",  (1, 0), (1,  0),  0),
    ]))
    return t


def _skills_2col_table(category_map: dict[str, list[str]], st: dict) -> Table:
    """
    Two-column layout where each column is a (Category | Skills) pair.
    Clean professional skills block matching the attached image format.
    """
    categories = list(category_map.items())   # [(label, [skill, ...]), ...]
    # Split into left and right halves
    mid   = (len(categories) + 1) // 2
    left  = categories[:mid]
    right = categories[mid:]

    # Pad right side so both columns are the same height
    while len(right) < len(left):
        right.append(("", []))

    COL_LABEL_W = 1.4 * inch
    COL_VALUE_W = (CONTENT_W / 2) - COL_LABEL_W - 0.15 * inch
    SPACER_W    = 0.3 * inch

    col_widths = [COL_LABEL_W, COL_VALUE_W, SPACER_W, COL_LABEL_W, COL_VALUE_W]

    rows = []
    style_cmds = [
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
        ("TOPPADDING",    (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        # No borders for clean look like the image
    ]

    for r_idx, (l_entry, r_entry) in enumerate(zip(left, right)):
        l_label, l_skills = l_entry
        r_label, r_skills = r_entry

        # Join skills with comma and space (clean, scannable — no bullets)
        l_val = ", ".join(l_skills)
        r_val = ", ".join(r_skills)

        rows.append([
            Paragraph(_md_to_xml(l_label), st["skill_label"]) if l_label else Paragraph("", st["skill_label"]),
            Paragraph(_md_to_xml(l_val), st["skill_value"]),
            Paragraph("", st["normal"]),   # spacer column
            Paragraph(_md_to_xml(r_label), st["skill_label"]) if r_label else Paragraph("", st["skill_label"]),
            Paragraph(_md_to_xml(r_val), st["skill_value"]),
        ])

    t = Table(rows, colWidths=col_widths, spaceBefore=4, spaceAfter=10)
    t.setStyle(TableStyle(style_cmds))
    return t


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → flowables
# ─────────────────────────────────────────────────────────────────────────────
def _parse_resume(md_text: str, st: dict) -> list:
    """
    Parses a markdown resume into ReportLab flowables.

    Header block (before first ##):
      # Candidate Name
      Job Title / Tagline          ← NEW: rendered in blue under name
      email | phone | location | linkedin

    Sections:
      ## SECTION NAME
      ### Role | Company | Date | Location
      - bullet
      **Category:** skill1, skill2   ← collected into 2-column skills table
    """
    story = []
    lines = [ln.rstrip() for ln in md_text.splitlines()]
    i = 0

    saw_name    = False
    saw_title   = False   # job-title line consumed
    saw_contact = False

    while i < len(lines):
        raw  = lines[i]
        line = raw.strip()

        # ── blank ─────────────────────────────────────────────────────────────
        if not line:
            if story:
                story.append(Spacer(1, 0.04 * inch))
            i += 1
            continue

        # ── # Name ───────────────────────────────────────────────────────────
        if line.startswith("# ") and not saw_name:
            story.append(Paragraph(line[2:].strip(), st["name"]))
            saw_name = True
            i += 1
            continue

        # ── Job title line (right after name, before contact) ─────────────────
        # Heuristic: no @, no digits-only phone pattern, no linkedin/github,
        # not a section header → treat as job title tagline
        if saw_name and not saw_title and not saw_contact and not line.startswith("#"):
            is_contact = (
                "@" in line
                or "linkedin" in line.lower()
                or "github" in line.lower()
                or "|" in line
                or re.search(r'\+?\d[\d\s\-().]{6,}', line)
            )
            if not is_contact:
                story.append(Paragraph(line, st["job_title_header"]))
                saw_title = True
                i += 1
                continue

        # ── Contact line ──────────────────────────────────────────────────────
        if saw_name and not saw_contact and not line.startswith("#"):
            contact = re.sub(r'\s*\|\s*', '  •  ', line).strip()
            story.append(Paragraph(contact, st["contact"]))
            story.append(_hr(thickness=1.2, spaceBefore=0, spaceAfter=10))
            saw_contact = True
            i += 1
            continue

        # ── ## Section header ─────────────────────────────────────────────────
        if line.startswith("## "):
            story.append(Paragraph(line[3:].strip().upper(), st["section"]))
            story.append(_hr(thickness=0.8, spaceBefore=1, spaceAfter=5))
            i += 1
            continue

        # ── ### Experience / Education entry ─────────────────────────────────
        if line.startswith("### "):
            parts    = [p.strip() for p in line[4:].split("|")]
            role     = parts[0] if len(parts) > 0 else ""
            company  = parts[1] if len(parts) > 1 else ""
            date     = parts[2] if len(parts) > 2 else ""
            location = parts[3] if len(parts) > 3 else ""
            story.append(_exp_row(role, company, date, location, st))
            i += 1
            continue

        # ── **Category:** val1, val2  → accumulate then emit 2-col table ─────
        if line.startswith("**") and ":**" in line:
            # Collect consecutive skill lines into an ordered dict
            category_map: dict[str, list[str]] = {}
            current_cat = None

            while i < len(lines):
                cur = lines[i].strip()

                # New category header
                if cur.startswith("**") and ":**" in cur:
                    idx         = cur.index(":**")
                    current_cat = cur[2:idx].strip()
                    value_part  = cur[idx + 3:].strip()
                    category_map[current_cat] = []
                    if value_part:
                        # Skills on same line as header
                        for s in re.split(r'[,;]', value_part):
                            s = s.strip()
                            if s:
                                category_map[current_cat].append(s)
                    i += 1

                # Bullet under current category
                elif cur.startswith(("- ", "* ", "• ")) and current_cat:
                    skill = cur[2:].strip()
                    if skill:
                        category_map[current_cat].append(skill)
                    i += 1

                # Comma-separated continuation lines (common after **Category:** on its own line)
                elif current_cat and cur and not cur.startswith("##") and not cur.startswith("###"):
                    for s in re.split(r'[,;]', cur):
                        s = s.strip()
                        if s:
                            category_map[current_cat].append(s)
                    i += 1

                elif not cur:
                    i += 1
                    continue

                else:
                    break

            if category_map:
                story.append(_skills_2col_table(category_map, st))
            continue

        # ── bullet ────────────────────────────────────────────────────────────
        if line.startswith(("- ", "* ", "• ")):
            story.append(
                Paragraph(f"• {_md_to_xml(line[2:].strip())}", st["bullet"])
            )
            i += 1
            continue

        # ── horizontal rule ───────────────────────────────────────────────────
        if line in ("---", "***", "___"):
            story.append(_hr(color=LIGHT_GRAY, thickness=0.4))
            i += 1
            continue

        # ── plain paragraph ───────────────────────────────────────────────────
        story.append(Paragraph(_md_to_xml(line), st["normal"]))
        i += 1

    return story


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────
def create_resume_pdf(markdown_content: str,
                      filename_title: str = "Resume") -> BytesIO:
    """
    Convert a markdown resume into a professional PDF.

    Expected markdown shape:
        # Full Name
        Job Title / Role Tagline
        email | phone | city | linkedin

        ## EXPERIENCE
        ### Role | Company | Jan 2024 – Present | City, ST
        - Bullet one
        - Bullet two

        ## SKILLS
        **Languages:** Python, Java, Kotlin
        **AI/ML:** TensorFlow, PyTorch
        - scikit-learn          ← bullets under a category also work

    Returns:
        BytesIO buffer — pass directly to st.download_button or write to file.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=L_MAR, rightMargin=R_MAR,
        topMargin=T_MAR,  bottomMargin=B_MAR,
        title=filename_title,
    )
    st = _styles()
    doc.build(_parse_resume(markdown_content, st))
    buffer.seek(0)
    return buffer


# ─────────────────────────────────────────────────────────────────────────────
# Word (.docx) generation
# ─────────────────────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str) -> None:
    """Fill a table cell background with a hex color (e.g. 'E0E7FF')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs) -> None:
    """Set borders on a table cell. kwargs keys: top/bottom/left/right/insideH/insideV."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, color in kwargs.items():
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   "single")
        tag.set(qn("w:sz"),    "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _add_hr(doc: "Document", color: str = "2C3E50", thickness: int = 4) -> None:
    """Add a horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


def _apply_bold_italic(run, text: str) -> None:
    """Handle **bold** and *italic* inline markdown in a single run."""
    # Simplified: strip markers and set flags
    if text.startswith("**") and text.endswith("**"):
        run.bold = True
        run.text = text[2:-2]
    elif text.startswith("*") and text.endswith("*"):
        run.italic = True
        run.text = text[1:-1]
    else:
        run.text = text


def _para_add_inline(para, text: str, bold: bool = False,
                     italic: bool = False, size: int = 10,
                     color: str | None = None) -> None:
    """Add a run to *para* with optional formatting."""
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def create_resume_docx(markdown_content: str) -> BytesIO:
    """Convert a markdown resume to a professionally styled .docx file."""
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin   = Inches(0.6)
        section.right_margin  = Inches(0.6)

    lines = [ln.rstrip() for ln in markdown_content.splitlines()]
    i = 0
    saw_name = saw_title = saw_contact = False

    while i < len(lines):
        raw  = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # ── # Name ───────────────────────────────────────────────────────────
        if line.startswith("# ") and not saw_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            _para_add_inline(p, line[2:].strip(), bold=True, size=22, color="0F172A")
            saw_name = True
            i += 1
            continue

        # ── Job title tagline ─────────────────────────────────────────────
        if saw_name and not saw_title and not saw_contact and not line.startswith("#"):
            is_contact = (
                "@" in line or "linkedin" in line.lower()
                or "github" in line.lower() or "|" in line
                or re.search(r'\+?\d[\d\s\-().]{6,}', line)
            )
            if not is_contact:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(3)
                _para_add_inline(p, line, size=12, color="3498DB")
                saw_title = True
                i += 1
                continue

        # ── Contact line ──────────────────────────────────────────────────
        if saw_name and not saw_contact and not line.startswith("#"):
            contact = re.sub(r'\s*\|\s*', '  •  ', line).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            _para_add_inline(p, contact, size=9, color="1E293B")
            _add_hr(doc, color="2C3E50", thickness=6)
            saw_contact = True
            i += 1
            continue

        # ── ## Section header ─────────────────────────────────────────────
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(2)
            _para_add_inline(p, line[3:].strip().upper(), bold=True, size=11, color="0F172A")
            _add_hr(doc, color="2C3E50", thickness=4)
            i += 1
            continue

        # ── ### Experience / Education entry ──────────────────────────────
        if line.startswith("### "):
            parts    = [p.strip() for p in line[4:].split("|")]
            role     = parts[0] if len(parts) > 0 else ""
            company  = parts[1] if len(parts) > 1 else ""
            date     = parts[2] if len(parts) > 2 else ""
            location = parts[3] if len(parts) > 3 else ""

            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            tbl.autofit = False
            tbl.columns[0].width = Inches(4.5)
            tbl.columns[1].width = Inches(2.8)
            left_cell  = tbl.rows[0].cells[0]
            right_cell = tbl.rows[0].cells[1]
            for cell in (left_cell, right_cell):
                for border in ("top", "bottom", "left", "right"):
                    _set_cell_border(cell, **{border: "FFFFFF"})

            lp = left_cell.paragraphs[0]
            lp.paragraph_format.space_before = Pt(6)
            left_text = role
            if company:
                left_text += f",  {company}"
            _para_add_inline(lp, role, bold=True, size=10, color="0F172A")
            if company:
                _para_add_inline(lp, f",  {company}", size=10, color="1E293B")

            rp = right_cell.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rp.paragraph_format.space_before = Pt(6)
            if date:
                _para_add_inline(rp, date, bold=True, size=9, color="475569")
            if location:
                _para_add_inline(rp, f"\n{location}", italic=True, size=9, color="475569")
            i += 1
            continue

        # ── **Category:** skills  → 2-column skills table ────────────────
        if line.startswith("**") and ":**" in line:
            category_map: dict[str, list[str]] = {}
            current_cat = None

            while i < len(lines):
                cur = lines[i].strip()
                if cur.startswith("**") and ":**" in cur:
                    idx         = cur.index(":**")
                    current_cat = cur[2:idx].strip()
                    value_part  = cur[idx + 3:].strip()
                    category_map[current_cat] = []
                    if value_part:
                        for s in re.split(r'[,;]', value_part):
                            s = s.strip()
                            if s:
                                category_map[current_cat].append(s)
                    i += 1
                elif cur.startswith(("- ", "* ", "• ")) and current_cat:
                    skill = cur[2:].strip()
                    if skill:
                        category_map[current_cat].append(skill)
                    i += 1
                elif current_cat and cur and not cur.startswith("##") and not cur.startswith("###"):
                    for s in re.split(r'[,;]', cur):
                        s = s.strip()
                        if s:
                            category_map[current_cat].append(s)
                    i += 1
                elif not cur:
                    i += 1
                    continue
                else:
                    break

            if not category_map:
                continue

            cats = list(category_map.items())
            mid   = (len(cats) + 1) // 2
            left_cats  = cats[:mid]
            right_cats = cats[mid:]
            while len(right_cats) < len(left_cats):
                right_cats.append(("", []))

            # 5-column table: label | value | spacer | label | value (clean, minimal — like print resume)
            tbl = doc.add_table(rows=len(left_cats), cols=5)
            tbl.style = "Table Grid"
            tbl.autofit = False
            col_widths = [Inches(1.35), Inches(2.1), Inches(0.15), Inches(1.35), Inches(2.1)]
            for ci, w in enumerate(col_widths):
                for row in tbl.rows:
                    row.cells[ci].width = w

            for ri, ((ll, ls), (rl, rs)) in enumerate(zip(left_cats, right_cats)):
                cells = tbl.rows[ri].cells
                for cell in cells:
                    for edge in ("top", "bottom", "left", "right"):
                        _set_cell_border(cell, **{edge: "FFFFFF"})

                def _fill(cell, label: str, skills: list[str], is_label: bool) -> None:
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after  = Pt(4)
                    _para_add_inline(p, label if is_label else ", ".join(skills),
                                     bold=is_label, size=9,
                                     color="0F172A" if is_label else "1E293B")

                _fill(cells[0], ll, ls, True)
                _fill(cells[1], ll, ls, False)
                _fill(cells[3], rl, rs, True)
                _fill(cells[4], rl, rs, False)

            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # ── bullet ────────────────────────────────────────────────────────
        if line.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent   = Inches(0.15)
            p.paragraph_format.space_before  = Pt(1)
            p.paragraph_format.space_after   = Pt(2)
            _para_add_inline(p, line[2:].strip(), size=10, color="1E293B")
            i += 1
            continue

        # ── plain paragraph ───────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _para_add_inline(p, line, size=10, color="1E293B")
        i += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_cover_letter_docx(cover_letter_content: str) -> BytesIO:
    """Convert a cover letter to a professionally styled .docx file."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(20)
    _para_add_inline(p, datetime.now().strftime("%B %d, %Y"), size=11, color="475569")

    in_sig = False
    for line in cover_letter_content.splitlines():
        s = line.strip()
        if not s:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue
        is_sig = any(s.lower().startswith(kw) for kw in
                     ("sincerely", "best regards", "regards", "thank you", "yours truly"))
        if is_sig:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            _para_add_inline(p, s, bold=True, size=11, color="0F172A")
            in_sig = True
        elif in_sig:
            p = doc.add_paragraph()
            _para_add_inline(p, s, bold=True, size=11, color="0F172A")
            in_sig = False
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            _para_add_inline(p, s, size=11, color="1E293B")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_cover_letter_pdf(cover_letter_content: str,
                             name: str = "") -> BytesIO:
    """Convert plain-text / lightly-markdown cover letter into a single-page PDF."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=1.0 * inch, rightMargin=1.0 * inch,  # Reduced margins for more space
        topMargin=0.8 * inch,  bottomMargin=0.8 * inch,   # Reduced margins for more space
        title="Cover Letter",
    )
    base = getSampleStyleSheet()

    date_style = ParagraphStyle(
        "CLDate", parent=base["Normal"],
        fontName="Helvetica", fontSize=10, leading=14,  # Smaller font
        textColor=LIGHT_GRAY, alignment=TA_RIGHT, spaceAfter=20,
    )
    body_style = ParagraphStyle(
        "CLBody", parent=base["Normal"],
        fontName="Helvetica", fontSize=10, leading=15,  # Smaller font and tighter spacing
        textColor=BODY_TEXT, spaceAfter=8,               # Reduced spacing
    )
    sig_style = ParagraphStyle(
        "CLSig", parent=base["Normal"],
        fontName="Helvetica-Bold", fontSize=10, leading=15,  # Smaller font
        textColor=DARK_GRAY, spaceBefore=20, spaceAfter=4,   # Reduced spacing
    )

    story = [Paragraph(datetime.now().strftime("%B %d, %Y"), date_style)]
    in_sig = False

    for line in cover_letter_content.splitlines():
        s = line.strip()
        if not s:
            story.append(Spacer(1, 0.05 * inch))  # Smaller spacer
            continue

        if any(s.lower().startswith(kw) for kw in
               ("sincerely", "best regards", "regards", "thank you", "yours truly")):
            story.append(Spacer(1, 0.2 * inch))  # Smaller spacer
            story.append(Paragraph(s, sig_style))
            in_sig = True
        elif in_sig:
            story.append(Paragraph(s, sig_style))
            in_sig = False
        else:
            story.append(Paragraph(_md_to_xml(s), body_style))

    if len(story) > 40:
        logger.warning(
            "Cover letter may overflow a single page (%d story elements). "
            "Consider shortening the content.",
            len(story),
        )

    doc.build(story)
    buffer.seek(0)
    return buffer
